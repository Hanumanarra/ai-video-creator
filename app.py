'''import os
import PyPDF2
from openai import OpenAI
from dotenv import load_dotenv, find_dotenv

_ = load_dotenv(find_dotenv())


api_key = os.environ.get('OPENAI_API_KEY')
if not api_key:
    print("Error: OPENAI_API_KEY not found in .env file or environment variables.")
    exit()

client = OpenAI(api_key=api_key)


model = 'gpt-4o-mini'
temperature = 0.3

max_tokens = 2000

book_text = ""  


file_path = "SQL Cheat Sheet Views Stored Procedures and Transactions.pdf"  
with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        if not reader.pages:
            print(f"Warning: The PDF '{file_path}' has no pages.")
        for page_num, page in enumerate(reader.pages):
            try:
                extracted_page_text = page.extract_text()  
                if extracted_page_text:
                    book_text += extracted_page_text + "\n" 
            except Exception as e:
                print(f"Error extracting text from page {page_num + 1}: {e}")



if not book_text.strip():
    print("Error: No text could be extracted from the PDF. The PDF might be image-based or empty.")
    exit()





user_task_description = "List out 1st page of content:"


user_prompt_content = f"{user_task_description}\n\n--- BOOK CONTENT ---\n{book_text}\n--- END BOOK CONTENT ---"


api_messages = [
  
    {"role": "user", "content": user_prompt_content}
]


def get_response_from_openai(messages_payload):
    try:
        completion = client.chat.completions.create(
            model=model,
            messages=messages_payload,
            temperature=temperature,
            max_tokens=max_tokens
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"Error calling OpenAI API: {e}"


print(f"Extracted text from '{file_path}' and sending to OpenAI...\n")
result = get_response_from_openai(api_messages)
print("\n--- OpenAI Response (10 Facts) ---")
print(result)
'''

'''
import os
import requests
from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
import openai
import json

load_dotenv()
app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)

# --- Configuration & Setup ---
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

HEYGEN_API_KEY = os.getenv('HEYGEN_API_KEY')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

try:
    # --- FIXED: Initialize the OpenAI client only ONCE ---
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
except Exception as e:
    print(f"Failed to initialize OpenAI client: {e}")
    client = None

if not HEYGEN_API_KEY: print("CRITICAL: HEYGEN_API_KEY is not loaded.")
if not OPENAI_API_KEY or not client: print("CRITICAL: OPENAI_API_KEY is not loaded or client failed.")


# --- NEW: Intelligent Voice Mapping ---
# Define the voice IDs you provided
FEMALE_VOICE_ID = '8f389c2237194f80b50fe7632dcc17b8'
MALE_VOICE_ID = '0ac81e725f4948dfa9638ceca216bcfa'

# Map specific avatar IDs from your script.js to their correct voice ID
# CRUCIAL: Keep this map in sync with your script.js `avatars` array
AVATAR_TO_VOICE_MAP = {
    # Female Avatars
    '20cf0b98ae164abdb4a814dab98e69ca_39260': FEMALE_VOICE_ID, # Mia
    'Abigail_expressive_2024112501': FEMALE_VOICE_ID,         # Abigail
    'Carlotta_Half_Front_public': FEMALE_VOICE_ID,            # Carlotta
    'Hada_Suit_Front_public': FEMALE_VOICE_ID,                # Hada
    
    # Male Avatars
    'Diran_iPad_Front_public': MALE_VOICE_ID                  # Diran
}
# --- End of New Section ---


def extract_text_from_file(filepath):
    text = ""
    try:
        if filepath.lower().endswith('.pdf'):
            with open(filepath, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    if page.extract_text():
                        text += page.extract_text() + "\n"
        elif filepath.lower().endswith('.docx'):
            doc = Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from {filepath}: {e}")
        raise ValueError(f"Could not read file: {os.path.basename(filepath)}")


@app.route("/")
def home():
    return render_template("index.html")

@app.route('/api/process-content', methods=['POST'])
def process_content():
    try:
        task = request.form.get('task', 'Extract key points for a presentation')
        num_points = request.form.get('num_points', 10)
        input_method = request.form.get('inputMethod')
        language = request.form.get('language', 'english').capitalize() # Get selected language
        source_text = ""

        if input_method == 'file':
            if 'file' not in request.files or not request.files['file'].filename:
                return jsonify({'error': 'No file selected'}), 400
            file=request.files.get('file')
            if file is None:
                raise ValueError("NO file is uploaded")
            #file = request.files['file']
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            source_text = extract_text_from_file(filepath)
        elif input_method == 'text':
            source_text = request.form.get('text_content', '')
        
        if not source_text.strip():
            return jsonify({'error': 'Empty content provided'}), 400

        # --- MODIFIED: Prompt now includes language instruction ---
        prompt = f"""
        Task: {task}.
        From the following content, please generate exactly {num_points} concise, well-written bullet points suitable for a video presentation script.
        Each bullet point should be on a new line. Do not include any introductory text like 'Here are the points:'.
        
        CRITICAL: The entire output script MUST be in the {language} language.

        Content:
        ---
        {source_text}
        ---
        """
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": f"You are an expert scriptwriter for short video presentations. You are fluent in many languages, including {language}."},
                {"role": "user", "content": prompt}
            ]
        )
        script = response.choices[0].message.content.strip()
        return jsonify({'script': script})

    except Exception as e:
        print(f"ERROR in /api/process-content: {str(e)}")
        return jsonify({"error": f"An internal server error occurred: {str(e)}"}), 500

@app.route('/api/generate-video', methods=['POST'])
def generate_video():
    if not HEYGEN_API_KEY:
        return jsonify({'error': 'HeyGen API key is not configured.'}), 500
    try:
        data = request.get_json()
        script = data.get("script")
        avatar_id = data.get("avatarId")
        
        if not script or not avatar_id:
            return jsonify({'error': 'Script and Avatar ID are required.'}), 400

        # --- MODIFIED: Automatically determine voice ID from the map ---
        # We use the female voice as a safe default if the avatar isn't in our map
        voice_id = AVATAR_TO_VOICE_MAP.get(avatar_id, FEMALE_VOICE_ID)
        
        headers = { "X-Api-Key": HEYGEN_API_KEY, "Content-Type": "application/json" }
        
        payload = {
            "video_inputs": [{
                "character": { "type": "avatar", "avatar_id": avatar_id, "avatar_style": "normal" },
                "voice": { "type": "text", "input_text": script, "voice_id": voice_id }
            }],
            "test": False,
            "dimension": { "width": 1920, "height": 1080 }
        }
        
        api_url = "https://api.heygen.com/v2/video/generate"
        response = requests.post(api_url, json=payload, headers=headers, timeout=30)
        response_data = response.json()

        print(f"HeyGen Response Status: {response.status_code}")
        print(f"HeyGen Response Body: {response_data}")

        if response.status_code == 200 and not response_data.get('error'):
            video_id = response_data.get('data', {}).get('video_id')
            return jsonify({"videoId": video_id})
        else:
            error_details = response_data.get('error', {}).get('message', 'Unknown API error.')
            return jsonify({ "error": "Failed to start video generation.", "details": error_details }), response.status_code

    except Exception as e:
        print(f"ERROR in /api/generate-video: {str(e)}")
        return jsonify({"error": "An internal server error occurred.", "details": str(e)}), 500


@app.route('/api/video-status/<string:video_id>', methods=['GET'])
def video_status(video_id):
    if not HEYGEN_API_KEY:
        return jsonify({'error': 'HeyGen API key is not configured.'}), 500
    try:
        status_url = f"https://api.heygen.com/v1/video_status.get?video_id={video_id}"
        headers = {"X-Api-Key": HEYGEN_API_KEY}
        
        response = requests.get(status_url, headers=headers, timeout=20)
        response.raise_for_status() 
        response_data = response.json()
        
        print(f"Polling V1 status for {video_id}: {response_data}")

        video_data_from_heygen = response_data.get('data', {})
        if not video_data_from_heygen:
            error_message = response_data.get('message', 'No data object in HeyGen response.')
            return jsonify({"error": error_message, "details": response_data}), 500
        
        # This is the fix: Wrap the data from HeyGen in our own 'data' key
        # to match what the JavaScript expects.
        return jsonify({'data': video_data_from_heygen})
            
    except requests.exceptions.HTTPError as http_err:
        status_code = http_err.response.status_code
        try:
             details = http_err.response.json()
        except json.JSONDecodeError:
             details = http_err.response.text
        print(f"HTTP ERROR in /api/video-status: {http_err} - Response: {details}")
        return jsonify({"error": "Failed to get video status from HeyGen.", "details": details}), status_code
    except Exception as e:
        print(f"ERROR in /api/video-status: {str(e)}")
        return jsonify({"error": "An internal server error occurred.", "details": str(e)}), 500
if __name__ == '__main__':
    app.run(debug=True)'''

import os
import requests
from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
import openai
import json

load_dotenv()
app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)

# --- Configuration & Setup ---
# UPLOAD_FOLDER is removed as it's not needed for in-memory processing

HEYGEN_API_KEY = os.getenv('HEYGEN_API_KEY')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

try:
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
except Exception as e:
    print(f"Failed to initialize OpenAI client: {e}")
    client = None

if not HEYGEN_API_KEY: print("CRITICAL: HEYGEN_API_KEY is not loaded.")
if not OPENAI_API_KEY or not client: print("CRITICAL: OPENAI_API_KEY is not loaded or client failed.")

# --- Intelligent Voice Mapping ---
FEMALE_VOICE_ID = '8f389c2237194f80b50fe7632dcc17b8'
MALE_VOICE_ID = '8f0944e10aad4e989bce8f76807b6f36'

# CRUCIAL: IDs here MUST MATCH the IDs in script.js
AVATAR_TO_VOICE_MAP = {
    # Female Avatars
    'Adriana_Business_Front_2_public': FEMALE_VOICE_ID,
    'Abigail_expressive_2024112501': FEMALE_VOICE_ID,
    'Carlotta_Business_Front_public': FEMALE_VOICE_ID, # Corrected ID
    'Annelise_public_3': FEMALE_VOICE_ID,
    # Male Avatars
    'Armando_Casual_Front_public': MALE_VOICE_ID
}

# ✅ FIXED: Function now accepts a file stream for Vercel compatibility
def extract_text_from_file(file_stream):
    """Extracts text from a file stream (for PDF or DOCX) in memory."""
    text = ""
    filename = secure_filename(file_stream.filename)
    try:
        if filename.lower().endswith('.pdf'):
            reader = PdfReader(file_stream)
            for page in reader.pages:
                if page.extract_text():
                    text += page.extract_text() + "\n"
        elif filename.lower().endswith('.docx'):
            doc = Document(file_stream)
            for para in doc.paragraphs:
                text += para.text + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        raise ValueError(f"Could not read file: {filename}")

@app.route("/")
def home():
    return render_template("index.html")

@app.route('/api/process-content', methods=['POST'])
def process_content():
    try:
        task = request.form.get('task', 'Extract key points for a presentation')
        num_points = request.form.get('num_points', 10)
        input_method = request.form.get('inputMethod')
        language = request.form.get('language', 'english').capitalize()
        source_text = ""

        if input_method == 'file':
            if 'file' not in request.files or not request.files['file'].filename:
                # ✅ FIXED: Better error handling
                return jsonify({'error': 'No file was selected for upload.'}), 400
            file = request.files['file']
            # ✅ FIXED: Process file in memory, do not save to disk
            source_text = extract_text_from_file(file)
        elif input_method == 'text':
            source_text = request.form.get('text_content', '')

        if not source_text.strip():
            return jsonify({'error': 'Empty content provided. Please upload a file or paste text.'}), 400

        prompt = f"""
        Task: {task}.
        From the following content, please generate exactly {num_points} concise, well-written bullet points suitable for a video presentation script.
        Each bullet point should be on a new line. Do not include any introductory text like 'Here are the points:'.
        CRITICAL: The entire output script MUST be in the {language} language.
        Content:
        ---
        {source_text}
        ---
        """
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": f"You are an expert scriptwriter for short video presentations. You are fluent in many languages, including {language}."},
                {"role": "user", "content": prompt}
            ]
        )
        script = response.choices[0].message.content.strip()
        
        # ✅ CRITICAL FIX: Return the script itself, not a generic message
        return jsonify({'script': script})

    except Exception as e:
        print(f"ERROR in /api/process-content: {str(e)}")
        # Check for authentication error specifically
        if "invalid_api_key" in str(e):
             return jsonify({"error": "Your OpenAI API key is invalid or has expired. Please check your .env file."}), 401
        return jsonify({"error": f"An internal server error occurred: {str(e)}"}), 500

# The rest of your app.py (generate_video, video_status) is correct and doesn't need changes.
# I've included them here for completeness.

@app.route('/api/generate-video', methods=['POST'])
def generate_video():
    if not HEYGEN_API_KEY:
        return jsonify({'error': 'HeyGen API key is not configured.'}), 500
    try:
        data = request.get_json()
        script = data.get("script")
        avatar_id = data.get("avatarId")
        
        if not script or not avatar_id:
            return jsonify({'error': 'Script and Avatar ID are required.'}), 400

        voice_id = AVATAR_TO_VOICE_MAP.get(avatar_id, FEMALE_VOICE_ID)
        
        headers = { "X-Api-Key": HEYGEN_API_KEY, "Content-Type": "application/json" }
        payload = {
            "video_inputs": [{
                "character": { "type": "avatar", "avatar_id": avatar_id, "avatar_style": "normal" },
                "voice": { "type": "text", "input_text": script, "voice_id": voice_id }
            }],
            "test": False,
            "dimension": { "width": 1920, "height": 1080 }
        }
        
        api_url = "https://api.heygen.com/v2/video/generate"
        response = requests.post(api_url, json=payload, headers=headers, timeout=30)
        response_data = response.json()

        if response.status_code == 200 and not response_data.get('error'):
            video_id = response_data.get('data', {}).get('video_id')
            return jsonify({"videoId": video_id})
        else:
            error_details = response_data.get('error', {}).get('message', 'Unknown API error.')
            return jsonify({ "error": "Failed to start video generation.", "details": error_details }), response.status_code

    except Exception as e:
        print(f"ERROR in /api/generate-video: {str(e)}")
        return jsonify({"error": "An internal server error occurred.", "details": str(e)}), 500


@app.route('/api/video-status/<string:video_id>', methods=['GET'])
def video_status(video_id):
    if not HEYGEN_API_KEY:
        return jsonify({'error': 'HeyGen API key is not configured.'}), 500
    try:
        status_url = f"https://api.heygen.com/v1/video_status.get?video_id={video_id}"
        headers = {"X-Api-Key": HEYGEN_API_KEY}
        
        response = requests.get(status_url, headers=headers, timeout=20)
        response.raise_for_status() 
        response_data = response.json()
        
        video_data_from_heygen = response_data.get('data', {})
        if not video_data_from_heygen:
            error_message = response_data.get('message', 'No data object in HeyGen response.')
            return jsonify({"error": error_message, "details": response_data}), 500
        
        return jsonify({'data': video_data_from_heygen})
            
    except requests.exceptions.HTTPError as http_err:
        status_code = http_err.response.status_code
        try:
             details = http_err.response.json()
        except json.JSONDecodeError:
             details = http_err.response.text
        return jsonify({"error": "Failed to get video status from HeyGen.", "details": details}), status_code
    except Exception as e:
        print(f"ERROR in /api/video-status: {str(e)}")
        return jsonify({"error": "An internal server error occurred.", "details": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)